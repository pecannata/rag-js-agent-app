#!/usr/bin/env python3
"""
Microsoft Word Document Text Extraction Script
Extracts text from .docx files using python-docx and handles various edge cases.
Includes OCR capability for extracting text from images.
"""

import sys
import json
import argparse
import io
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print(json.dumps({
        "success": False,
        "error": "python-docx not installed. Please run: pip install python-docx",
        "text": "",
        "paragraphCount": 0
    }))
    sys.exit(1)

# OCR dependencies - optional
OCR_AVAILABLE = False
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    pass  # OCR will be skipped if dependencies not available

def extract_text_from_images_in_doc(doc):
    """
    Extract text from images embedded in Word document using OCR.
    
    Args:
        doc: python-docx Document object
        
    Returns:
        str: Extracted text from all images in the document
    """
    if not OCR_AVAILABLE:
        print("OCR not available - skipping image text extraction", file=sys.stderr)
        return ""
    
    ocr_text = ""
    image_count = 0
    
    try:
        # Access the document's relationships to find embedded images
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_count += 1
                    print(f"Processing image {image_count}: {rel.target_ref}", file=sys.stderr)
                    
                    # Get the image data
                    image_part = rel.target_part
                    image_data = image_part.blob
                    
                    # Convert to PIL Image
                    image_stream = io.BytesIO(image_data)
                    image = Image.open(image_stream)
                    
                    print(f"Image {image_count} size: {image.size}, mode: {image.mode}", file=sys.stderr)
                    
                    # Perform OCR with multiple configurations for better results
                    ocr_configs = [
                        '--psm 6',   # Uniform block of text
                        '--psm 8',   # Single word
                        '--psm 7',   # Single text line
                        '--psm 11',  # Sparse text
                        '--psm 13'   # Raw line
                    ]
                    
                    best_text = ""
                    for config in ocr_configs:
                        try:
                            extracted_text = pytesseract.image_to_string(image, config=config).strip()
                            if len(extracted_text) > len(best_text):
                                best_text = extracted_text
                        except:
                            continue
                    
                    if best_text:
                        ocr_text += f"\n[OCR from Image {image_count}]\n{best_text}\n"
                        print(f"OCR extracted from image {image_count}: {best_text[:100]}...", file=sys.stderr)
                    else:
                        print(f"No text found in image {image_count}", file=sys.stderr)
                        
                except Exception as e:
                    print(f"Error processing image {image_count}: {e}", file=sys.stderr)
                    continue
    
    except Exception as e:
        print(f"Error accessing document images: {e}", file=sys.stderr)
    
    if image_count > 0:
        print(f"Processed {image_count} images for OCR", file=sys.stderr)
    else:
        print("No images found in document", file=sys.stderr)
    
    return ocr_text

def extract_text_from_docx(docx_path, max_paragraphs=1000):
    """
    Extract text from Word document file.
    
    Args:
        docx_path (str): Path to the .docx file
        max_paragraphs (int): Maximum number of paragraphs to process
        
    Returns:
        dict: Result containing success status, text, and metadata
    """
    try:
        doc = Document(docx_path)
        
        print(f"OCR Available: {OCR_AVAILABLE}", file=sys.stderr)
        
        extracted_text = ""
        paragraph_texts = []
        processed_paragraphs = 0
        
        # Extract text from embedded images using OCR
        ocr_text = ""
        if OCR_AVAILABLE:
            print("Extracting text from images using OCR...", file=sys.stderr)
            ocr_text = extract_text_from_images_in_doc(doc)
        
        for i, paragraph in enumerate(doc.paragraphs):
            if i >= max_paragraphs:
                break
                
            try:
                para_text = paragraph.text.strip()
                
                if para_text:  # Only include non-empty paragraphs
                    paragraph_texts.append({
                        "paragraph": i + 1,
                        "text": para_text
                    })
                    extracted_text += f"{para_text}\n\n"
                    processed_paragraphs += 1
                    
            except Exception as para_error:
                paragraph_texts.append({
                    "paragraph": i + 1,
                    "text": f"[Error extracting text from paragraph {i + 1}: {str(para_error)}]"
                })
                extracted_text += f"[Error extracting paragraph {i + 1}]\n\n"
        
        # Also extract text from tables if present
        table_text = ""
        table_count = 0
        
        for table in doc.tables:
            table_count += 1
            table_text += f"\n--- Table {table_count} ---\n"
            
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_text += f"{row_text}\n"
            
            table_text += "\n"
        
        # Combine paragraph and table text
        if table_text.strip():
            extracted_text += f"\n{table_text}"
        
        # Add OCR text from images if any was found
        if ocr_text.strip():
            extracted_text += f"\n{ocr_text}"
        
        # Clean up the text
        cleaned_text = extracted_text.strip()
        if cleaned_text:
            # Normalize excessive whitespace but preserve paragraph breaks
            lines = []
            for line in cleaned_text.split('\n'):
                cleaned_line = line.strip()
                if cleaned_line or (lines and lines[-1]):  # Keep empty lines only if preceded by content
                    lines.append(cleaned_line)
            cleaned_text = '\n'.join(lines)
        
        return {
            "success": True,
            "text": cleaned_text,
            "paragraphCount": len(doc.paragraphs),
            "processedParagraphs": processed_paragraphs,
            "tableCount": table_count,
            "paragraphTexts": paragraph_texts,
            "hasText": bool(cleaned_text.strip()),
            "ocrAvailable": OCR_AVAILABLE,
            "ocrTextFound": bool(ocr_text.strip())
        }
        
    except FileNotFoundError:
        return {
            "success": False,
            "error": f"Word document not found: {docx_path}",
            "text": "",
            "paragraphCount": 0
        }
    except Exception as e:
        error_msg = str(e)
        if "not a zip file" in error_msg.lower():
            return {
                "success": False,
                "error": "Invalid Word document format. Please ensure the file is a valid .docx file (not .doc).",
                "text": "",
                "paragraphCount": 0
            }
        elif "permission" in error_msg.lower():
            return {
                "success": False,
                "error": f"Permission denied accessing file: {docx_path}",
                "text": "",
                "paragraphCount": 0
            }
        else:
            return {
                "success": False,
                "error": f"Unexpected error processing Word document: {error_msg}",
                "text": "",
                "paragraphCount": 0
            }

def main():
    parser = argparse.ArgumentParser(description='Extract text from Microsoft Word (.docx) files')
    parser.add_argument('docx_path', help='Path to the .docx file')
    parser.add_argument('--max-paragraphs', type=int, default=1000, help='Maximum paragraphs to process')
    
    args = parser.parse_args()
    
    result = extract_text_from_docx(args.docx_path, args.max_paragraphs)
    print(json.dumps(result, indent=2))

if __name__ == "__main__":
    main()
